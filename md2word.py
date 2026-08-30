import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


DEFAULT_MD_PATH = r"C:\Users\Gary\Desktop\jpg2word\jpg_files\final_book.md"
DEFAULT_DOCX_PATH = r"C:\Users\Gary\Desktop\jpg2word\jpg_files\final_book.docx"


def set_run_font(run, size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_font(paragraph, size: int | None = None, bold: bool | None = None) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold)


def set_cell_text(cell, text: str) -> None:
    cell.text = text.strip()
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_run_font(run, size=10)


def apply_document_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    normal = document.styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in [("Heading 1", 18), ("Heading 2", 15), ("Heading 3", 13)]:
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True


def apply_section_layout(section) -> None:
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def set_image_name_footer(section, image_name: str) -> None:
    section.footer.is_linked_to_previous = False
    paragraph = section.footer.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(image_name)
    set_run_font(run, size=10)


def is_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2


def is_separator_row(line: str) -> bool:
    cells = parse_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def parse_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


def add_table(document: Document, table_lines: list[str]) -> None:
    rows = [parse_table_row(line) for line in table_lines if not is_separator_row(line)]
    rows = [row for row in rows if row]
    if not rows:
        return

    col_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.autofit = True

    for row_index, row in enumerate(rows):
        for col_index in range(col_count):
            text = row[col_index] if col_index < len(row) else ""
            set_cell_text(table.cell(row_index, col_index), text)

    document.add_paragraph()


def add_markdown_paragraph(document: Document, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        return

    heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
    if heading_match:
        level = min(len(heading_match.group(1)), 3)
        paragraph = document.add_heading(heading_match.group(2).strip(), level=level)
        set_paragraph_font(paragraph, bold=True)
        return

    paragraph = document.add_paragraph()
    run = paragraph.add_run(stripped)
    set_run_font(run, size=11)


def split_markdown_pages(lines: list[str]) -> list[list[str]]:
    pages: list[list[str]] = []
    current: list[str] = []

    for line in lines:
        if line.strip() == "---":
            if any(item.strip() for item in current):
                pages.append(current)
            current = []
            continue
        current.append(line)

    if any(item.strip() for item in current):
        pages.append(current)
    return pages


def load_image_names(manifest_path: Path | None, pages: list[list[str]]) -> list[str]:
    names: list[str] = []
    if manifest_path and manifest_path.exists():
        for line in manifest_path.read_text(encoding="utf-8").splitlines():
            match = re.match(r"\s*\d+\.\s+`([^`]+\.(?:jpg|jpeg))`\s+->", line, flags=re.IGNORECASE)
            if match:
                names.append(match.group(1))

    if len(names) < len(pages):
        for page in pages[len(names) :]:
            text = "\n".join(page)
            match = re.search(r"\b([^\\/:*?\"<>|`\s]+\.(?:jpg|jpeg))\b", text, flags=re.IGNORECASE)
            names.append(match.group(1) if match else f"page_{len(names) + 1}")

    return names[: len(pages)]


def add_page_content(document: Document, page_lines: list[str]) -> None:
    table_buffer: list[str] = []

    def flush_table() -> None:
        if table_buffer:
            add_table(document, table_buffer)
            table_buffer.clear()

    for line in page_lines:
        stripped = line.strip()
        if is_table_line(stripped):
            table_buffer.append(stripped)
            continue

        flush_table()
        add_markdown_paragraph(document, line)

    flush_table()


def markdown_to_docx(md_path: str | Path, docx_path: str | Path, manifest_path: str | Path | None = None) -> Path:
    md_path = Path(md_path)
    docx_path = Path(docx_path)
    if not md_path.exists():
        raise FileNotFoundError(f"Markdown 文件不存在：{md_path}")

    manifest = Path(manifest_path) if manifest_path else md_path.with_name("manifest.md")

    document = Document()
    apply_document_styles(document)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    pages = split_markdown_pages(lines)
    image_names = load_image_names(manifest, pages)

    for index, page_lines in enumerate(pages):
        section = document.sections[-1] if index == 0 else document.add_section(WD_SECTION.NEW_PAGE)
        apply_section_layout(section)
        set_image_name_footer(section, image_names[index])
        add_page_content(document, page_lines)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(docx_path)
    return docx_path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="将 final_book.md 转换成 Word docx。")
    parser.add_argument("--md", default=DEFAULT_MD_PATH, help="输入 Markdown 文件。")
    parser.add_argument("--docx", default=DEFAULT_DOCX_PATH, help="输出 Word 文件。")
    parser.add_argument("--manifest", default="", help="图片清单 manifest.md，默认使用 Markdown 同目录下的 manifest.md。")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    manifest_path = args.manifest or None
    docx_path = markdown_to_docx(args.md, args.docx, manifest_path)
    print(f"已生成 Word：{docx_path}", flush=True)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
